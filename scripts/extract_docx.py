#!/usr/bin/env python3
"""Extract headings, paragraphs, and tables from a DOCX into JSON or text."""

import argparse
import json
import re
from pathlib import Path

from docx import Document


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def extract(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = []
    headings = []
    for idx, para in enumerate(doc.paragraphs, 1):
        text = clean(para.text)
        if not text:
            continue
        style = para.style.name if para.style is not None else ""
        item = {"index": idx, "style": style, "text": text}
        paragraphs.append(item)
        if style.lower().startswith("heading") or re.match(r"^(\d+[\.\、]|第[一二三四五六七八九十]+[章节])", text):
            headings.append(item)

    tables = []
    for t_idx, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append({"index": t_idx, "rows": rows})

    return {
        "source": str(path),
        "paragraph_count": len(paragraphs),
        "heading_count": len(headings),
        "table_count": len(tables),
        "char_count": sum(len(p["text"]) for p in paragraphs),
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--format", choices=["json", "text"], default="json")
    parser.add_argument("--out")
    args = parser.parse_args()

    data = extract(Path(args.docx))
    if args.format == "json":
        output = json.dumps(data, ensure_ascii=False, indent=2)
    else:
        lines = [f"# SOURCE: {data['source']}"]
        lines.append("## HEADINGS")
        for h in data["headings"]:
            lines.append(f"[P{h['index']:04d}][{h['style']}] {h['text']}")
        lines.append("## PARAGRAPHS")
        for p in data["paragraphs"]:
            lines.append(f"[P{p['index']:04d}][{p['style']}] {p['text']}")
        lines.append("## TABLES")
        for t in data["tables"]:
            lines.append(f"[TABLE {t['index']}]")
            for i, row in enumerate(t["rows"], 1):
                lines.append(f"R{i}: " + " | ".join(row))
        output = "\n".join(lines)

    if args.out:
        Path(args.out).write_text(output, encoding="utf-8")
    else:
        print(output)


if __name__ == "__main__":
    main()

