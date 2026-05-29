#!/usr/bin/env python3
"""Convert a confirmed protocol Markdown file into a styled DOCX.

This script is intentionally conservative: it preserves Markdown content and
basic structure instead of asking the model to recreate protocol text.
"""

import argparse
import re
from pathlib import Path

from docx import Document

try:
    from style_unify_docx import apply_styles
except ImportError:  # pragma: no cover - fallback for unusual execution paths
    apply_styles = None


CHINESE_NUMERAL = "一二三四五六七八九十百千万"


def strip_front_matter(lines):
    if not lines or lines[0].strip() != "---":
        return lines
    for idx in range(1, len(lines)):
        if lines[idx].strip() == "---":
            return lines[idx + 1 :]
    return lines


def is_table_separator(line):
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and "|" in lines[idx].strip():
        line = lines[idx].strip()
        if line:
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        idx += 1
    if len(rows) >= 2 and is_table_separator(lines[start + 1]):
        return [rows[0]] + rows[2:], idx
    return None, start


def add_inline_runs(paragraph, text, bold=False):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run_bold = bold
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            run_bold = True
        elif part.startswith("`") and part.endswith("`"):
            part = part[1:-1]
        paragraph.add_run(part).bold = run_bold


def add_paragraph(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    add_inline_runs(para, text)
    return para


def add_table(doc, rows):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            add_inline_runs(para, text, bold=(r_idx == 0))


def chinese_heading_level(line):
    text = line.strip()
    if re.match(rf"^[{CHINESE_NUMERAL}]+[、.．]\s*", text):
        return 1
    if re.match(rf"^（[{CHINESE_NUMERAL}]+）\s*", text):
        return 2
    if re.match(r"^\d+(\.\d+)*[、.．]\s+", text):
        return 3
    return None


def markdown_to_docx(md_path, out_path):
    md_path = Path(md_path)
    out_path = Path(out_path)
    lines = strip_front_matter(md_path.read_text(encoding="utf-8").splitlines())

    doc = Document()
    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped in {"---", "***", "___"}:
            doc.add_page_break()
            idx += 1
            continue

        table_rows, next_idx = parse_table(lines, idx)
        if table_rows:
            add_table(doc, table_rows)
            idx = next_idx
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            idx += 1
            continue

        chinese_level = chinese_heading_level(stripped)
        if chinese_level and len(stripped) <= 80:
            doc.add_heading(stripped, level=chinese_level)
            idx += 1
            continue

        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet_match:
            add_paragraph(doc, bullet_match.group(1), style="List Bullet")
            idx += 1
            continue

        number_match = re.match(r"^\d+[.)、]\s+(.+)$", stripped)
        if number_match:
            add_paragraph(doc, number_match.group(1), style="List Number")
            idx += 1
            continue

        quote_match = re.match(r"^>\s?(.+)$", stripped)
        if quote_match:
            add_paragraph(doc, quote_match.group(1), style="Intense Quote")
            idx += 1
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if image_match:
            alt, target = image_match.groups()
            add_paragraph(doc, f"[图片: {alt or target}]")
            idx += 1
            continue

        add_paragraph(doc, stripped)
        idx += 1

    if apply_styles:
        apply_styles(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(out_path)


def main():
    parser = argparse.ArgumentParser(
        description="Convert a confirmed protocol Markdown source file to DOCX."
    )
    parser.add_argument("input_md", help="Confirmed protocol Markdown file")
    parser.add_argument("--out", required=True, help="Output DOCX path")
    args = parser.parse_args()
    markdown_to_docx(args.input_md, args.out)


if __name__ == "__main__":
    main()
