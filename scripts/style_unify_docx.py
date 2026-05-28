#!/usr/bin/env python3
"""Apply simple, conservative clinical-protocol Word styles to a DOCX."""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_font(run, east_asia="宋体", latin="Times New Roman", size=10.5, bold=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_styles(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = "Times New Roman"
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            st.font.size = Pt(size)
            st.font.bold = True
            st.paragraph_format.space_before = Pt(12)
            st.paragraph_format.space_after = Pt(6)

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("[待确认]") or "[TBD]" in text or "[需用户确认]" in text:
            for run in para.runs:
                run.font.highlight_color = 7  # yellow
        for run in para.runs:
            if para.style and para.style.name.startswith("Heading"):
                set_font(run, east_asia="黑体", size=run.font.size.pt if run.font.size else 12, bold=True)
            else:
                set_font(run)

    for table in doc.tables:
        if table.rows:
            for cell in table.rows[0].cells:
                shade_cell(cell, "D9EAD3")
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_font(run, east_asia="黑体", size=10.5, bold=True)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_font(run, size=10)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    doc = Document(args.input)
    apply_styles(doc)
    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(args.out)


if __name__ == "__main__":
    main()

